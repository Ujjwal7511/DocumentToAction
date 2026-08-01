"""Document ingestion: parsing (.txt / .md / .docx) and type detection."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import BinaryIO, Optional, Union

from llm_client import LLMError, generate_json, is_configured

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".txt", ".md", ".markdown", ".docx", ".pdf"}
MAX_DOCUMENTS = 3

DOC_TYPES = [
    "meeting_notes",
    "requirement_draft",
    "implementation_notes",
    "project_update",
    "decision_record",
]

DOC_TYPE_SCHEMA: dict = {
    "type": "object",
    "properties": {
        "doc_type": {
            "type": "string",
            "enum": DOC_TYPES + ["unknown"],
        },
        "confidence": {"type": "number"},
        "rationale": {"type": "string"},
    },
    "required": ["doc_type", "confidence", "rationale"],
}


def parse_txt(data: Union[bytes, str]) -> str:
    """Decode plain-text / markdown bytes or return a string as-is."""
    if isinstance(data, str):
        return data
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_docx(data: bytes) -> str:
    """Extract text from a .docx file, preserving paragraph breaks."""
    from io import BytesIO

    from docx import Document as DocxDocument

    document = DocxDocument(BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def parse_pdf(data: bytes) -> str:
    """Extract text from a PDF using a lightweight fallback."""
    try:
        import pypdf
    except Exception:  # noqa: BLE001
        raise ValueError("PyPDF is required for PDF support. Install it with pip install pypdf")

    from io import BytesIO

    reader = pypdf.PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page for page in pages if page).strip()


def parse_file(
    filename: str,
    data: Union[bytes, BinaryIO, str],
) -> str:
    """Parse an uploaded file into plain text.

    Raises ValueError for unsupported extensions or empty content after parse.
    """
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if hasattr(data, "read"):
        raw = data.read()
    else:
        raw = data

    if ext == ".docx":
        if isinstance(raw, str):
            raw = raw.encode("utf-8")
        text = parse_docx(raw)
    elif ext == ".pdf":
        if isinstance(raw, str):
            raw = raw.encode("utf-8")
        text = parse_pdf(raw)
    else:
        text = parse_txt(raw)

    text = text.strip()
    return text


def split_into_sections(text: str) -> list[dict]:
    """Split document text into sections keyed by heading or paragraph index.

    Returns a list of dicts:
      {heading, start, end, text}
    Used for source linking of extracted items.
    """
    if not text or not text.strip():
        return []

    # Markdown / ADR-style headings
    heading_re = re.compile(r"^(#{1,6}\s+.+|[A-Z][A-Za-z0-9 /&\-]{2,60}:?\s*)$", re.MULTILINE)
    matches = list(heading_re.finditer(text))

    sections: list[dict] = []
    if not matches:
        # Fall back to blank-line paragraphs
        offset = 0
        parts = re.split(r"\n\s*\n", text)
        for i, part in enumerate(parts):
            part_stripped = part.strip()
            if not part_stripped:
                continue
            start = text.find(part_stripped, offset)
            if start < 0:
                start = offset
            end = start + len(part_stripped)
            sections.append(
                {
                    "heading": f"Paragraph {i + 1}",
                    "start": start,
                    "end": end,
                    "text": part_stripped,
                }
            )
            offset = end
        return sections

    # Content before first heading
    if matches[0].start() > 0:
        preamble = text[: matches[0].start()].strip()
        if preamble:
            sections.append(
                {
                    "heading": "Preamble",
                    "start": 0,
                    "end": matches[0].start(),
                    "text": preamble,
                }
            )

    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        heading = match.group(0).strip().lstrip("#").strip()
        body = text[start:end].strip()
        sections.append(
            {
                "heading": heading or f"Section {i + 1}",
                "start": start,
                "end": end,
                "text": body,
            }
        )
    return sections


def _heuristic_doc_type(text: str, filename: str) -> tuple[str, float, str]:
    """Keyword-based fallback when LLM is unavailable."""
    lower = (text + " " + filename).lower()
    scores = {
        "meeting_notes": sum(
            k in lower for k in ("agenda", "attendees", "standup", "minutes", "action items", "meeting")
        ),
        "requirement_draft": sum(
            k in lower for k in ("requirement", "shall", "user story", "acceptance criteria", "must ")
        ),
        "implementation_notes": sum(
            k in lower for k in ("implementation", "api", "schema", "deploy", "refactor", "code")
        ),
        "project_update": sum(
            k in lower for k in ("status", "progress", "blocker", "this week", "update", "milestone")
        ),
        "decision_record": sum(
            k in lower for k in ("decision", "adr", "we decided", "alternatives", "consequences")
        ),
    }
    best = max(scores, key=scores.get)
    if scores[best] == 0:
        return "unknown", 0.2, "No distinctive keywords found; classified as unknown."
    confidence = min(0.4 + 0.1 * scores[best], 0.75)
    return best, confidence, f"Heuristic keyword score={scores[best]} for {best}."


def detect_document_type(text: str, filename: str = "") -> dict:
    """Identify document type via LLM (preferred) or heuristics.

    Returns dict with keys: doc_type, confidence, rationale.
    """
    if not text or not text.strip():
        return {
            "doc_type": "unknown",
            "confidence": 0.0,
            "rationale": "Document is empty; type cannot be determined.",
        }

    if not is_configured():
        doc_type, confidence, rationale = _heuristic_doc_type(text, filename)
        return {
            "doc_type": doc_type,
            "confidence": confidence,
            "rationale": rationale + " (LLM unavailable — heuristic used.)",
        }

    excerpt = text[:6000]
    prompt = (
        "Classify the following project document into exactly one type:\n"
        f"{', '.join(DOC_TYPES)}, or unknown.\n\n"
        f"Filename: {filename}\n\n"
        f"Document:\n---\n{excerpt}\n---\n\n"
        "Return doc_type, confidence (0-1), and a short rationale."
    )
    try:
        result = generate_json(prompt, DOC_TYPE_SCHEMA)
        doc_type = str(result.get("doc_type", "unknown"))
        if doc_type not in DOC_TYPES and doc_type != "unknown":
            doc_type = "unknown"
        confidence = float(result.get("confidence", 0.5))
        confidence = max(0.0, min(1.0, confidence))
        rationale = str(result.get("rationale", ""))
        return {"doc_type": doc_type, "confidence": confidence, "rationale": rationale}
    except (LLMError, Exception) as exc:  # noqa: BLE001
        logger.warning("LLM type detection failed, using heuristic: %s", exc)
        doc_type, confidence, rationale = _heuristic_doc_type(text, filename)
        return {
            "doc_type": doc_type,
            "confidence": confidence,
            "rationale": f"{rationale} (LLM error: {exc})",
        }


def ingest_upload(
    filename: str,
    data: Union[bytes, BinaryIO, str],
) -> dict:
    """Parse a file and detect its type. Returns a ready-to-persist dict."""
    content = parse_file(filename, data)
    sections = split_into_sections(content) if content else []
    type_info = detect_document_type(content, filename)
    return {
        "filename": filename,
        "content": content,
        "sections": sections,
        "doc_type": type_info["doc_type"],
        "doc_type_confidence": type_info["confidence"],
        "doc_type_rationale": type_info["rationale"],
        "is_empty": not bool(content.strip()) if content else True,
    }


def validate_upload_batch(filenames: list[str]) -> Optional[str]:
    """Return an error message if the upload batch is invalid, else None."""
    if not filenames:
        return "Please upload at least one document."
    if len(filenames) > MAX_DOCUMENTS:
        return f"You may upload at most {MAX_DOCUMENTS} documents."
    for name in filenames:
        ext = Path(name).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            return f"Unsupported file type for '{name}'. Use .txt, .md, .docx, or .pdf."
    return None
